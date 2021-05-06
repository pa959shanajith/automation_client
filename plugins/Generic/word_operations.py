#from bs4 import UnicodeDammit
import PyPDF2
import string
import json
from lxml import etree
import docx
from pprint import pprint
import xmltodict
import generic_constants
from constants import *
import os
import subprocess
import logging
log = logging.getLogger('word_operations.py')
import logger
import platform
if SYSTEM_OS == 'Windows':
    import fitz

class WordFile():
    def writeWordFile(self, filename, filetext, operation, file_encoding = None, *args):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = OUTPUT_CONSTANT
        err_msg = None
        try:
            fileName,file_ext=os.path.splitext(filename)
            if file_ext.lower() ==".docx":
                if not os.path.exists(os.path.normpath(filename)):
                    log.debug("Write File- File Not Found. Creating a new File with name : "+filename)
                if file_encoding:
                    filetext = filetext.encode(file_encoding)
                if operation.lower() == "overwrite":
                    file = open(filename,"w")
                    file.write(filetext)
                    file.close()
                elif operation.lower()=="append":
                    file = open(filename,'a+')
                    file.write(filetext)
                    file.close()
                status = TEST_RESULT_PASS
                result = TEST_RESULT_TRUE
            else:
                err_msg=generic_constants.INVALID_INPUT
                log.debug('Invalid File Extension')
        except Exception as e:
            log.error(e)
            err_msg = generic_constants.INVALID_INPUT
        if err_msg != None:
            log.error(err_msg)
            logger.print_on_console(err_msg)
        ##logger.info("Write File => " + ERROR_CODE_DICT['MSG_STATUS'] + status)
        ##logger.print_on_console("Write File => " + ERROR_CODE_DICT['MSG_STATUS'] + status)
        return  status, result, output, err_msg

    def readWorddoc(self,filename,parano,*args): ##parano = -1 -> for all paragraphs
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = OUTPUT_CONSTANT
        err_msg=None
        paras = []
        para = None

        try:
            num = int(parano)
            if os.path.exists(os.path.normpath(filename)):
                doc=docx.Document(filename)
                for para in doc.paragraphs:
                    if(para.text and not(para.text.isspace())):
                        paras.append(para.text)
                if(parano==-1):
                    output = "\n".join(paras)
                else:
                    output = paras[num-1]
                log.debug('Total Number of Paragraphs :'+str(len(paras)))
                status = TEST_RESULT_PASS
                result = TEST_RESULT_TRUE
            else:
                err_msg = generic_constants.FILE_NOT_EXISTS
                log.error(err_msg)
                logger.print_on_console("Read Paragraph From Word => Error : File Not Found")
        except Exception as e:
            if(isinstance(e, ValueError)):
                err_msg = ERROR_CODE_DICT['ERR_INVALID_INPUT']
            else:
                err_msg = ERROR_CODE_DICT['ERR_EXCEPTION']
            log.error(e) #@Delete
            logger.print_on_console(err_msg)
        return status, result, output, err_msg

    def readallcheckbox(self, filename):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = OUTPUT_CONSTANT
        err_msg = None

        try:
            if os.path.exists(os.path.normpath(filename)):
                doc = docx.Document(filename)
                from docx.oxml.ns import qn
                from lxml import etree
                major_number = 0
                main_answer = {}
                doc_ele = doc._element
                checkBoxes = doc_ele.xpath('.//w:checkBox')
                if not checkBoxes:
                    #w14 type
                    checkBoxes = doc_ele.xpath("//*[starts-with(local-name(), 'checkbox')]")
                    if not checkBoxes:
                        checkBoxes = doc_ele.xpath("//*[starts-with(local-name(), 'checkBox')]")
                if checkBoxes:
                    for checkBox in checkBoxes:
                        check_dict = {}
                        checkboxstatus = False
                        parent_itr = checkBox
                        nextstring = None
                        while(parent_itr.getparent() is not None and nextstring is None):
                            parent_itr = parent_itr.getparent()
                            sibling_obj = parent_itr
                            while(sibling_obj.getnext() is not None and nextstring is None):
                                if sibling_obj.tag.split('}')[-1] == "r":
                                    for t_tag_obj in sibling_obj.iter(qn("w:t")) :
                                        if t_tag_obj.text:
                                            nextstring = t_tag_obj.text
                                        break
                                sibling_obj = sibling_obj.getnext()

                        for ch in checkBox:
                            if "checked" in ch.tag:
                                if ch.get(qn('w:val')) is None:
                                    # not using qn(w14:val) because it gives w14 namespace undefined
                                    if ch.get('val') is None:
                                        checkboxstatus = True
                                    elif ch.get('val') == "1" or ch.get('val').lower() == "true":
                                        checkboxstatus = True
                                elif ch.get(qn('w:val')) == "1" or ch.get(qn('w:val')).lower() == "true":
                                    checkboxstatus = True
                        if checkboxstatus:
                            check_dict["status"] = "checked"
                        else:
                            check_dict["status"] = "unchecked"
                        if nextstring is None:
                            nextstring = ""
                        check_dict["nexttext"] = nextstring
                        major_number = major_number + 1
                        main_answer[major_number] = check_dict
                main_answer['count'] = str(major_number)
                output = json.dumps(main_answer)
                status = TEST_RESULT_PASS
                result = TEST_RESULT_TRUE
            else:
                err_msg = ERROR_CODE_DICT['ERR_FILE_NOT_FOUND_EXCEPTION']
                log.error(err_msg)
                logger.print_on_console("readallcheckbox => Error : File not Found")
        except Exception as e:
            if isinstance(e, ValueError):
                err_msg = ERROR_CODE_DICT['ERR_INVALID_INPUT']
                ##err_json = {"ERR_CODE": err_code, "ERR_MSG": ERROR_CODE_DICT['ERR_INVALID_INPUT']}
            else:
                err_code = ERROR_CODE_DICT['ERR_EXCEPTION']
                ##err_json = {"ERR_CODE": err_code, "ERR_MSG": ERROR_CODE_DICT['ERR_EXCEPTION'], "STACK_TRACE": str(e)}
            log.error(err_msg)
            logger.print_on_console(err_msg)
            log.error(e)#@Delete
        ##logger.info("Read All Checkbox => " + ERROR_CODE_DICT['MSG_STATUS'] + status)
        ##logger.print_on_console("Read All Checkbox => " + ERROR_CODE_DICT['MSG_STATUS'] + status)
        return status, result, output, err_msg

    def getAllTablesFromDoc(self, filename, *args):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = None
        err_msg = None
        try:
            if SYSTEM_OS == "Windows":
                import pythoncom
                import win32com.client as win32
                fileName,file_ext=os.path.splitext(filename)
                if file_ext.lower() ==".docx":
                    if os.path.exists(os.path.normpath(filename)):
                        pythoncom.CoInitialize()
                        doc=docx.Document(filename)
                        for t in range(len(doc.tables)):
                            table = doc.tables[t]
                            data = []
                            keys = None
                            for i, row in enumerate(table.rows):
                                text = (cell.text for cell in row.cells)
                                if i == 0:
                                    keys = tuple(text)
                                    continue
                                row_data = dict(zip(keys, text))
                                data.append(row_data)
                        output = data
                        status = TEST_RESULT_PASS
                        result = TEST_RESULT_TRUE
                    else:
                        err_msg = ERROR_CODE_DICT['ERR_FILE_NOT_FOUND_EXCEPTION']
                        logger.print_on_console("Error : File not Found")
                        log.error(err_msg)
                else:
                    err_msg=generic_constants.INVALID_INPUT
                    log.debug('Invalid File Extension')
        except Exception as e:
            err_msg = ERROR_CODE_DICT['ERR_EXCEPTION']
            log.error(e)
        if err_msg != None:
            log.error(err_msg)
            logger.print_on_console(err_msg)
        return status, result, output, err_msg

    def readjson(self, filename, jsonpara, *args):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = None
        err_msg = None
        try:
            fileName,file_ext=os.path.splitext(filename)
            if file_ext.lower() ==".json":
                if os.path.exists(os.path.normpath(filename)):
                    log.debug("Filename : " + filename)
                    vals= list(args)
                    jsonparameters= vals.insert(0,jsonpara)
                    divlist = [e if e.isdigit() else e for e in vals]
                    with open(filename) as data_file:
                        data = json.load(data_file)
                    for e in divlist:
                        if(e.isdigit()):
                            data = data[int(e)-1]
                        else:
                            data = data[''+e]
                    output = data
                    status = TEST_RESULT_PASS
                    result = TEST_RESULT_TRUE
                else:
                    err_msg = ERROR_CODE_DICT['ERR_FILE_NOT_FOUND_EXCEPTION']
                    log.error(err_msg)
                    logger.print_on_console("Error : File not Found")
            else:
                err_msg=generic_constants.INVALID_INPUT
                log.debug('Invalid File Extension')
        except Exception as e:
            if isinstance(e, ValueError):
                err_msg = ERROR_CODE_DICT['ERR_INVALID_INPUT']
            elif isinstance(e, KeyError):
                err_msg = ERROR_CODE_DICT['ERR_OBJECT_NOT_EXISTS']
            else:
                err_msg = ERROR_CODE_DICT['ERR_EXCEPTION']
            log.error(err_msg)
            log.error(e) #@delete
            logger.print_on_console(err_msg)
        log.info(ERROR_CODE_DICT['MSG_STATUS'] + status)
        return status, result, output, err_msg

    def readxml(self,filename,xmlpara,*args):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = None
        err_msg = None
        try:
            fileName,file_ext=os.path.splitext(filename)
            if file_ext.lower() ==".xml":
                if os.path.exists(os.path.normpath(filename)):
                    vals= list(args)
                    xmlparameters= vals.insert(0,xmlpara)
                    divlist = [e if e.isdigit() else e for e in vals]
                    contents = open(filename).read()
                    #Since this is being parsed in JSON, the order of xml is lost
                    para1 = xmltodict.parse(contents)
                    para2 = json.dumps(para1)
                    data = json.loads(para2)
                    for e in divlist:
                        if e.isdigit():
                            data = data[int(e)-1]
                        else:
                            data = data[''+e]
                    output = data
                    status = TEST_RESULT_PASS
                    result = TEST_RESULT_TRUE
                else:
                    err_msg = ERROR_CODE_DICT['ERR_FILE_NOT_FOUND_EXCEPTION']
                    log.error(err_msg)
                    logger.print_on_console("Error: File not Found")
            else:
                err_msg=generic_constants.INVALID_INPUT
                log.debug('Invalid File Extension')
        except Exception as e:
            if isinstance(e, ValueError):
                err_msg = ERROR_CODE_DICT['ERR_INVALID_INPUT']
            elif isinstance(e, KeyError):
                err_msg = ERROR_CODE_DICT['ERR_OBJECT_NOT_EXISTS']
            else:
                err_msg = ERROR_CODE_DICT['ERR_EXCEPTION']
            log.error(e)
            logger.print_on_console(err_msg)
        return status, result, output, err_msg

    def readPdf(self, filename,pageno, *args):
        status = TEST_RESULT_FAIL
        result = TEST_RESULT_FALSE
        output = None
        err_msg = None
        try:
            if os.path.exists(os.path.normpath(filename)):
                fileName,file_ext=os.path.splitext(filename)
                if file_ext.lower() ==".pdf":
                    pdfData=fitz.open(filename)
                    pagenumber=int(pageno)-1
                    if pagenumber<pdfData.pageCount:
                        page = pdfData[pagenumber]
                    else:
                        page = pdfData[0]
                    output=page.getText()
                    output=output.replace("\n"," ").replace('\r',' ')
                    status = TEST_RESULT_PASS
                    result = TEST_RESULT_TRUE
                else:
                    err_msg = 'Invalid Input'
                    logger.print_on_console('Invalid Input')
            else:
                logger.print_on_console("Read PDF: Error : File not Found")
        except Exception as e:
            if isinstance(e, ValueError):
                err_msg = ERROR_CODE_DICT['ERR_INVALID_INPUT']
            else:
                err_msg = ERROR_CODE_DICT['ERR_EXCEPTION']
            log.error(err_msg)
            log.error(e)
            logger.print_on_console(err_msg)
        return status, result, output, err_msg


